# Библиотека для работы с документами .docx
import docx
# Библотека для работы с ОС
import os
# Библотека для работы с датой
import datetime
# Импортируем собсвтенный модуль для работы с файлами
import file_management

# Получаем текущий год
current_year = datetime.datetime.now().year


class Tab46_2_Years:
    """Класс для анализа текущего и прошлого года"""
    def __init__(self):
        file_path_1 = input(f'Введите путь к файлу {current_year} года: ')
        file_path_2 = input(f'Введите путь к файлу {current_year - 1} года: ')

        successfully = False
        while successfully == False:
            try:
                self.file_1 = file_management.File(name='Tab46.docx', path=file_path_1)
                self.file_2 = file_management.File(name='Tab46.docx', path=file_path_2)
                self.file_processing()
                successfully = True
            except TypeError:
                # РАБОТАЕТ ДЛЯ ВСЕХ ОШИБОК - ПОЙМАТЬ НУЖНУЮ НАДО НАСТРОИТЬЬЬЬФ!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
                print('Введен неверный адрес к файлу')
                file_path_1 = input(f'Введите путь к файлу {current_year} года: ')
                file_path_2 = input(f'Введите путь к файлу {current_year - 1} года: ')

    def file_processing(self):
        """Обработка файла статистики"""

        # Создаем документ, в котором будут сохранены результаты анализа
        document = docx.Document()

        document.add_paragraph('Анализ краж велосипедов')
        paragraph_1 = document.add_paragraph()
        paragraph_1.add_run(
            f"На территории Брянской области {self.file_1.document.paragraphs[2].text} "
            f"зарегистрировано {self.file_1.document.tables[0].rows[11].cells[1].text} "
            f"(АППГ {self.file_2.document.tables[0].rows[11].cells[1].text}, "
            f"{self.file_1.document.tables[0].rows[11].cells[2].text}%) кражи велосипедов, "
            f"раскрыто {self.file_1.document.tables[0].rows[11].cells[3].text} "
            f"(АППГ {self.file_2.document.tables[0].rows[11].cells[3].text}, "
            f"{self.file_1.document.tables[0].rows[11].cells[4].text}%) краж, "
            f"приостановлено {self.file_1.document.tables[0].rows[11].cells[5].text} "
            f"(АППГ {self.file_2.document.tables[0].rows[11].cells[5].text}, "
            f"{self.file_1.document.tables[0].rows[11].cells[6].text}%) уголовных дел. "
            # раскрываемость = раскрыто / (раскрыто + приостановлено) * 100
            f"Раскрываемость составила "
            f"{round(int(self.file_1.document.tables[0].rows[11].cells[3].text) / (int(self.file_1.document.tables[0].rows[11].cells[3].text) + int(self.file_1.document.tables[0].rows[11].cells[5].text)) * 100, 1)}% "
            f"(АППГ {round(int(self.file_2.document.tables[0].rows[11].cells[3].text) / (int(self.file_2.document.tables[0].rows[11].cells[3].text) + int(self.file_2.document.tables[0].rows[11].cells[5].text)) * 100, 1)}%). "
        )

        # вызываем создание графика

        paragraph_2 = document.add_paragraph()
        paragraph_2.add_run(
            f"Время совершения краж велосипедов составляет: с 9 до 18 часов - "
            f"{self.file_1.document.tables[0].rows[11].cells[7].text} "
            f"(АППГ {self.file_2.document.tables[0].rows[11].cells[7].text}), с 19 до 24 часов - "
            f"{self.file_1.document.tables[0].rows[11].cells[8].text} "
            f"(АППГ {self.file_2.document.tables[0].rows[11].cells[8].text}), с 0 до 09 часов - "
            f"{self.file_1.document.tables[0].rows[11].cells[9].text} "
            f"(АППГ {self.file_2.document.tables[0].rows[11].cells[9].text})."
        )

        paragraph_3 = document.add_paragraph()
        paragraph_3.add_run(
            f"На территориях городов и поселков городского типа совершено "
            f"{self.file_1.document.tables[0].rows[11].cells[11].text} "
            f"(АППГ {self.file_2.document.tables[0].rows[11].cells[11].text}) краж велосипедов, "
            f"на территориях сельских поселений: {self.file_1.document.tables[0].rows[11].cells[10].text} "
            f"(АППГ {self.file_2.document.tables[0].rows[11].cells[10].text})."
        )

        # Сохраняем полученные результаты в файл

        successfully = False
        while successfully == False:
            """Выполняется пока файл не сохранится"""
            try:
                document.save('Анализ файла Tab46.docx')
                successfully = True
            except PermissionError:
                print('Закройте ранее созданный файл "Анализ файла Tab46.docx"')
                input('После закрытия старого файла нажмите клавишу ENTER: ')